import fitz  # PyMuPDF
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog, messagebox, Text, ttk
import re
import datetime

def save_start_date(filename):
    start_date = datetime.datetime.now()
    with open(filename, 'w') as file:
        file.write(start_date.strftime('%Y-%m-%d %H:%M:%S'))

def load_start_date(filename):
    with open(filename, 'r') as file:
        start_date_str = file.read()
        return datetime.datetime.strptime(start_date_str, '%Y-%m-%d %H:%M:%S')

def is_trial_period_over(start_date, trial_days=30):
    current_date = datetime.datetime.now()
    return (current_date - start_date).days > trial_days

def initialize_trial_period(filename='start_date.txt'):
    if not os.path.exists(filename):
        save_start_date(filename)
    start_date = load_start_date(filename)
    if is_trial_period_over(start_date):
        messagebox.showerror("Período de Teste Expirado", "Seu período de teste de 30 dias expirou.")
        return False
    return True

# Funções principais do programa

def extract_info_from_pdf(pdf_path):
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"No such file: '{pdf_path}'")
    
    doc = fitz.open(pdf_path)
    extracted_text = ""
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        extracted_text += page.get_text()
    
    return extracted_text

def filter_info(text):
    patterns = {
        "Veículo": re.compile(r'Modelo\s+([\w\/\s\-]+)'),
        "Placa": re.compile(r'Placa\s+(\w+)'),
        "Ano": re.compile(r'Ano de Fabricação\s+(\d{4})'),
        "Cor": re.compile(r'Cor\s+([\w]+)'),
        "Lote": re.compile(r'Lote\s+(\d+)'),
        "Nº motor": re.compile(r'Nº Motor\s+([\w\-]+)'),
        "Câmbio": re.compile(r'Nº Câmbio\s+([\w\-]+)'),
        "Chassi": re.compile(r'Chassi\s+([\w]+)'),
        "Potência e cc": re.compile(r'Potência \(cv\)\s+(\d+)\s+Cilindradas \(cc\)\s+(\d+)'),
        "Etiqueta": re.compile(r'Etiqueta\s+(\w+)')
    }

    filtered_info = {}
    
    for key, pattern in patterns.items():
        match = pattern.search(text)
        if match:
            if key == "Potência e cc":
                filtered_info[key] = f"{match.group(1)} cv, {match.group(2)} cc"
            else:
                filtered_info[key] = match.group(1)
        else:
            filtered_info[key] = "Informação não encontrada"

    return "\n".join([f"{k}: {v}" for k, v in filtered_info.items()]), filtered_info.get("Veículo", "documento_resultante")

def clean_filename(filename):
    invalid_chars = r'[<>:"/\|?*]'
    cleaned_filename = re.sub(invalid_chars, '_', filename)
    return cleaned_filename.strip()

def save_to_word(text, filename):
    cleaned_filename = clean_filename(filename)
    doc = Document()
    doc.add_heading('Informações Extraídas do PDF', 0)
    
    for line in text.split('\n'):
        parts = line.split(": ")
        if len(parts) == 2:
            key, value = parts
            doc.add_heading(key, level=1)
            doc.add_paragraph(value)
    
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos do Word", "*.docx")])
    if save_path:
        doc.save(save_path)

def process_pdf():
    try:
        pdf_path = filedialog.askopenfilename(filetypes=[("Arquivos PDF", "*.pdf")])
        if not pdf_path:
            raise ValueError("Nenhum arquivo PDF selecionado.")

        progress.start()
        extracted_text = extract_info_from_pdf(pdf_path)
        filtered_text, vehicle_name = filter_info(extracted_text)
        progress.stop()
        
        text_display.delete(1.0, tk.END)
        text_display.insert(tk.END, filtered_text)
        
        save_button.config(state=tk.NORMAL)
        save_button.vehicle_name = vehicle_name  # Armazena o nome do veículo para uso posterior
    except Exception as e:
        progress.stop()
        messagebox.showerror("Erro", f"Ocorreu um erro: {e}")

def show_help():
    messagebox.showinfo("Ajuda", "1. Clique em 'Selecionar PDF' para escolher o arquivo PDF.\n"
                                  "\n2. O texto extraído será exibido na área de texto.\n"
                                  "\n3. Na área de texto, você pode inserir e conferir as informações, formatando como quiser.\n"
                                  "\n4. Clique em 'Salvar como Word' para salvar o resultado em um arquivo Word.\n"
                                  "\n5. Ao salvar, tome cuidado com barras invertidas e outros caracteres não aceitos, salve corretamente o nome do seu arquivo.")

def show_feedback():
    messagebox.showinfo("Feedback", "O programa está na fase teste, logo, pode conter erros.\n"
                                    "\nCaso queira dar um feedback para continuarmos trabalhando juntos, envie um e-mail para: \n"
                                    "\nvictorjunqueira.prog@gmail.com")

def main():
    global text_display, save_button, progress

    # Verifica o período de teste antes de iniciar a interface gráfica
    if not initialize_trial_period():
        return

    root = tk.Tk()
    root.title("Extrator de Informações de PDF")

    menu_bar = tk.Menu(root)

    # Menu Ajuda
    help_menu = tk.Menu(menu_bar, tearoff=0)
    help_menu.add_command(label="Ajuda", command=show_help)
    menu_bar.add_cascade(label="Ajuda", menu=help_menu)

    # Menu Feedback
    feedback_menu = tk.Menu(menu_bar, tearoff=0)
    feedback_menu.add_command(label="Feedback", command=show_feedback)
    menu_bar.add_cascade(label="Feedback", menu=feedback_menu)

    root.config(menu=menu_bar)

    frame = tk.Frame(root)
    frame.pack(pady=10)

    select_button = tk.Button(frame, text="Selecionar PDF", command=process_pdf)
    select_button.grid(row=0, column=0, padx=10)

    save_button = tk.Button(frame, text="Salvar como Word", command=lambda: save_to_word(text_display.get(1.0, tk.END), save_button.vehicle_name), state=tk.DISABLED)
    save_button.grid(row=0, column=1, padx=10)

    progress = ttk.Progressbar(root, orient=tk.HORIZONTAL, length=400, mode='indeterminate')
    progress.pack(padx=10, pady=10)

    text_display = Text(root, wrap='word', width=80, height=20)
    text_display.pack(padx=10, pady=10)
    
    root.mainloop()

if __name__ == "__main__":
    main()